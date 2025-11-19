"""
导出服务
"""

import logging
import os
import tempfile
from typing import List, Dict, Optional
from fastapi import BackgroundTasks, HTTPException
from fastapi.responses import FileResponse
from zipfile import ZipFile, ZIP_DEFLATED
from io import BytesIO
from lxml import etree
from docx import Document

from database import get_collection
from utils.mongo import style_collection, check_result_collection
from utils.exceptions import (
    HTTP_400_BAD_REQUEST,
    HTTP_404_NOT_FOUND,
    HTTP_500_INTERNAL_SERVER_ERROR,
)
from utils.docx_utils import cleanup_temp_file

logger = logging.getLogger(__name__)


def merge_styles_xml(styles_list: List[Dict[str, Optional[str]]]) -> Optional[str]:
    """
    合并多个表格的样式XML
    
    Args:
        styles_list: 样式字典列表，每个字典包含 paragraph_style, table_style, run_style, cell_style
        
    Returns:
        合并后的完整 styles.xml 字符串
    """
    try:
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # 创建新的 styles 根元素
        styles_root = etree.Element(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styles',
            nsmap=namespaces
        )
        
        # 用于去重的集合，避免重复添加相同样式
        added_style_ids = set()
        
        # 遍历所有样式字典
        for styles_dict in styles_list:
            if not styles_dict:
                continue
                
            # 遍历每个样式类型
            for style_key, style_xml in styles_dict.items():
                if not style_xml:
                    continue
                    
                try:
                    # 解析样式XML
                    style_element = etree.fromstring(style_xml.encode('utf-8'))
                    
                    # 获取样式ID用于去重
                    style_id = style_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId')
                    
                    # 如果样式ID已存在，跳过（避免重复）
                    if style_id and style_id in added_style_ids:
                        continue
                    
                    # 添加到根元素
                    styles_root.append(style_element)
                    if style_id:
                        added_style_ids.add(style_id)
                        
                except Exception as e:
                    logger.warning(f"解析样式XML失败 ({style_key}): {e}")
                    continue
        
        # 如果没有样式，至少创建一个基本的styles.xml结构
        if len(styles_root) == 0:
            # 创建一个默认段落样式
            default_style = etree.SubElement(
                styles_root,
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}style',
                attrib={
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type': 'paragraph',
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId': 'Normal',
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}default': '1'
                }
            )
            name_elem = etree.SubElement(
                default_style,
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name',
                attrib={
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val': 'Normal'
                }
            )
        
        # 转换为字符串
        styles_xml_str = etree.tostring(
            styles_root,
            encoding='utf-8',
            xml_declaration=True,
            pretty_print=True
        ).decode('utf-8')
        
        return styles_xml_str
        
    except Exception as e:
        logger.error(f"合并styles.xml失败: {e}", exc_info=True)
        return None


def create_page_break() -> str:
    """
    创建Word下一页分节符XML
    
    Returns:
        下一页分节符的XML字符串
    """
    return '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:sectPr><w:type w:val="nextPage"/></w:sectPr></w:pPr></w:p>'


def add_table_center_alignment(elem: etree.Element) -> None:
    """
    为表格元素添加居中对齐属性
    如果表格的tblPr中没有tblAlign，则添加<w:tblAlign w:val="center"/>
    
    Args:
        elem: XML元素（可能是表格或其他元素）
    """
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # 如果是表格元素，处理它
    if elem.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl':
        # 查找或创建tblPr元素
        tbl_pr = elem.find('w:tblPr', namespaces)
        if tbl_pr is None:
            # 如果tblPr不存在，创建它（作为第一个子元素）
            tbl_pr = etree.SubElement(
                elem,
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr'
            )
        
        # 检查是否已经有tblAlign
        tbl_align = tbl_pr.find('w:tblAlign', namespaces)
        if tbl_align is None:
            # 如果没有tblAlign，添加它
            etree.SubElement(
                tbl_pr,
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblAlign',
                attrib={
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val': 'center'
                }
            )
    
    # 递归处理所有子元素
    for child in elem:
        add_table_center_alignment(child)


def is_empty_paragraph(elem: etree.Element) -> bool:
    """
    检查段落是否为空（只包含格式信息，没有实际内容）
    
    Args:
        elem: XML元素（段落）
        
    Returns:
        如果段落为空返回True，否则返回False
    """
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # 如果不是段落元素，返回False
    if elem.tag != '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p':
        return False
    
    # 检查是否有文本内容（包括run中的文本）
    runs = elem.findall('w:r', namespaces)
    for run in runs:
        # 检查run中的文本
        texts = run.findall('w:t', namespaces)
        for text_elem in texts:
            if text_elem.text and text_elem.text.strip():
                return False
        
        # 检查run中是否有其他内容（如图片、对象等）
        for child in run:
            if child.tag not in [
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t',
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr',
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br',
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tab',
            ]:
                return False
    
    # 检查段落本身的文本
    if elem.text and elem.text.strip():
        return False
    
    # 检查是否有非格式的子元素（如表格、图片等）
    for child in elem:
        if child.tag not in [
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr',
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r',
        ]:
            return False
    
    # 检查tail文本
    if elem.tail and elem.tail.strip():
        return False
    
    return True


def is_section_break_paragraph(elem: etree.Element) -> bool:
    """
    检查段落是否包含分节符
    
    Args:
        elem: XML元素（段落）
        
    Returns:
        如果段落包含分节符返回True，否则返回False
    """
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    if elem.tag != '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p':
        return False
    
    # 检查段落属性中是否有sectPr
    p_pr = elem.find('w:pPr', namespaces)
    if p_pr is not None:
        sect_pr = p_pr.find('w:sectPr', namespaces)
        if sect_pr is not None:
            return True
    
    return False


def remove_empty_paragraphs_after_section_breaks(elements: List[etree.Element]) -> List[etree.Element]:
    """
    删除分节符段落后面的空白段落（模拟按Delete键）
    
    Args:
        elements: 元素列表
        
    Returns:
        清理后的元素列表
    """
    if not elements:
        return elements
    
    result = []
    i = 0
    
    while i < len(elements):
        elem = elements[i]
        result.append(elem)
        
        # 如果当前元素是分节符段落
        if is_section_break_paragraph(elem):
            # 检查后面的元素是否是空白段落
            j = i + 1
            while j < len(elements):
                next_elem = elements[j]
                # 如果是空白段落，跳过（不添加到result）
                if is_empty_paragraph(next_elem):
                    j += 1
                # 如果是另一个分节符段落，也跳过（避免重复分节符）
                elif is_section_break_paragraph(next_elem):
                    j += 1
                else:
                    # 遇到非空白段落，停止删除
                    break
            
            # 更新索引
            i = j
        else:
            i += 1
    
    return result


def merge_document_xml(document_xml_list: List[str], insert_page_breaks: bool = True) -> str:
    """
    合并多个document.xml片段，中间插入下一页分节符
    
    Args:
        document_xml_list: document.xml字符串列表
        insert_page_breaks: 是否在表格之间插入下一页分节符
        
    Returns:
        合并后的document.xml字符串
    """
    if not document_xml_list:
        return ""
    
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    parser = etree.XMLParser(remove_blank_text=False)
    
    # 解析下一页分节符
    page_break_xml = create_page_break()
    page_break_elem = etree.fromstring(page_break_xml.encode('utf-8'), parser=parser)
    
    # 收集所有要合并的元素
    all_elements = []
    
    for i, doc_xml in enumerate(document_xml_list):
        if not doc_xml:
            continue
            
        try:
            # 尝试解析XML（可能包含多个根元素）
            try:
                # 尝试作为单个XML文档解析
                content_tree = etree.fromstring(doc_xml.encode('utf-8'), parser=parser)
                elements = [content_tree]
            except etree.XMLSyntaxError:
                # 如果包含多个根元素，包装在临时根元素中
                wrapped_xml = f'<root xmlns:w="{namespaces["w"]}">{doc_xml}</root>'
                wrapped_tree = etree.fromstring(wrapped_xml.encode('utf-8'), parser=parser)
                elements = list(wrapped_tree)
            
            # 添加所有元素，并处理表格对齐
            for elem in elements:
                # 为所有表格添加居中对齐
                add_table_center_alignment(elem)
                all_elements.append(elem)
            
            # 如果不是最后一个，插入下一页分节符
            if insert_page_breaks and i < len(document_xml_list) - 1:
                all_elements.append(page_break_elem)
                
        except Exception as e:
            logger.warning(f"解析document.xml失败: {e}")
            continue
    
    # 如果没有元素，返回空字符串
    if not all_elements:
        return ""
    
    # 删除分节符段落后面的空白段落（模拟按Delete键）
    all_elements = remove_empty_paragraphs_after_section_breaks(all_elements)
    
    # 转换为字符串（只返回元素内容，不包含body标签）
    result_parts = []
    for elem in all_elements:
        result_parts.append(etree.tostring(elem, encoding='unicode', method='xml'))
    
    return ''.join(result_parts)


async def export_docx(
    table_ids: List[str], background_tasks: BackgroundTasks
) -> FileResponse:
    """
    导出表格为DOCX文件
    支持导出多个相同table_id的表格

    Args:
        table_ids: 表格ID列表（允许重复）
        background_tasks: 后台任务

    Returns:
        FileResponse对象
    """
    temp_file = None
    try:
        logger.info(f"收到文件下载请求: table_ids_count={len(table_ids)}")

        collection = get_collection()

        # 检查依赖是否初始化成功
        if collection is None:
            logger.error("文件下载失败: ChromaDB未初始化")
            raise HTTPException(
                status_code=HTTP_500_INTERNAL_SERVER_ERROR,
                detail="系统初始化失败，请检查ChromaDB",
            )

        # 参数验证
        if not table_ids or len(table_ids) == 0:
            logger.warning("文件下载失败: 表格ID数组为空")
            raise HTTPException(
                status_code=HTTP_400_BAD_REQUEST, detail="表格ID数组不能为空"
            )

        # 使用table_id集合进行批量查询（去重以提高查询效率）
        unique_table_ids = list(set(table_ids))
        logger.info(f"去重后的唯一table_id数量: {len(unique_table_ids)}")

        # 从ChromaDB批量查询表格数据
        try:
            results = collection.get(ids=unique_table_ids)
            logger.info(f"ChromaDB查询成功: 请求{len(unique_table_ids)}个唯一表格ID")
        except Exception as e:
            logger.error(f"ChromaDB查询失败: error={e}", exc_info=True)
            raise HTTPException(
                status_code=HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"ChromaDB查询失败: {str(e)}",
            )

        # 检查查询结果
        if not results or "ids" not in results or len(results["ids"]) == 0:
            logger.warning(f"文件下载失败: 未找到指定的表格数据 table_ids={unique_table_ids}")
            raise HTTPException(
                status_code=HTTP_404_NOT_FOUND, detail="未找到指定的表格数据"
            )

        # 构建table_id到数据的映射
        id_to_data = {}
        ids = results["ids"]
        documents = results.get("documents", [])
        metadatas = results.get("metadatas", [])

        for i, table_id in enumerate(ids):
            id_to_data[table_id] = {
                "document": documents[i] if i < len(documents) else "",
                "metadata": metadatas[i] if i < len(metadatas) else {},
            }

        # 从MongoDB批量查询样式信息
        styles_map = {}
        if style_collection is not None:
            try:
                # 使用table_id集合查询MongoDB
                mongo_results = style_collection.find({"table_id": {"$in": unique_table_ids}})
                for doc in mongo_results:
                    table_id = doc.get("table_id")
                    styles = doc.get("styles", {})
                    if table_id and styles:
                        styles_map[table_id] = styles
                logger.info(f"MongoDB查询成功: 找到{len(styles_map)}个表格的样式信息")
            except Exception as e:
                logger.warning(f"MongoDB查询失败: {e}，将继续使用空样式")
        else:
            logger.warning("MongoDB collection未初始化，跳过样式信息查询")

        # 按照请求顺序收集document.xml和styles
        document_xml_list = []
        styles_list = []

        for table_id in table_ids:
            if table_id not in id_to_data:
                logger.warning(f"表格ID未找到: table_id={table_id}")
                continue

            data = id_to_data[table_id]
            table_xml = data["document"]

            if not table_xml:
                logger.warning(f"表格XML为空: table_id={table_id}")
                continue

            document_xml_list.append(table_xml)
            
            # 获取对应的样式信息
            styles = styles_map.get(table_id, {})
            styles_list.append(styles)

        # 检查是否成功收集到表格
        if len(document_xml_list) == 0:
            logger.error(f"文件下载失败: 未能成功收集任何表格 table_ids={table_ids}")
            raise HTTPException(
                status_code=HTTP_400_BAD_REQUEST,
                detail="未能成功收集任何表格，请检查表格ID和数据",
            )

        logger.info(f"成功收集{len(document_xml_list)}个表格")

        # 合并所有表格的styles_xml
        merged_styles_xml = merge_styles_xml(styles_list)
        if not merged_styles_xml:
            logger.warning("合并styles.xml失败，将使用默认样式")

        # 合并所有表格的document.xml，中间插入下一页分节符
        merged_document_xml = merge_document_xml(document_xml_list, insert_page_breaks=True)
        if not merged_document_xml:
            logger.error("合并document.xml失败")
            raise HTTPException(
                status_code=HTTP_500_INTERNAL_SERVER_ERROR,
                detail="合并document.xml失败",
            )

        # 创建临时docx文件
        try:
            with tempfile.NamedTemporaryFile(
                mode="wb", suffix=".docx", delete=False
            ) as tmp:
                temp_file = tmp.name
        except Exception as e:
            logger.error(f"创建临时文件失败: {e}", exc_info=True)
            raise HTTPException(
                status_code=HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"创建临时文件失败: {str(e)}",
            )

        # 创建新的空DOCX文档作为模板
        doc = Document()
        template_buffer = BytesIO()
        doc.save(template_buffer)
        template_buffer.seek(0)

        # 读取模板DOCX并修改
        try:
            with ZipFile(template_buffer, 'r') as template_zip:
                # 创建新的ZIP文件
                output_buffer = BytesIO()
                with ZipFile(output_buffer, 'w', ZIP_DEFLATED) as new_zip:
                    # 复制模板中的所有文件（跳过document.xml和styles.xml，后面会写入）
                    files_to_skip = {'word/document.xml', 'word/styles.xml'}
                    for item in template_zip.infolist():
                        if item.filename not in files_to_skip:
                            data = template_zip.read(item.filename)
                            new_zip.writestr(item.filename, data)

                    # 读取原始document.xml作为模板
                    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    doc_xml = template_zip.read('word/document.xml')
                    doc_tree = etree.fromstring(doc_xml)

                    # 解析合并后的document.xml内容
                    parser = etree.XMLParser(remove_blank_text=False)
                    try:
                        # 尝试解析为单个XML文档
                        content_tree = etree.fromstring(merged_document_xml.encode('utf-8'), parser=parser)
                        elements_to_add = [content_tree]
                    except etree.XMLSyntaxError:
                        # 如果包含多个根元素，包装在临时根元素中
                        wrapped_xml = f'<root xmlns:w="{namespaces["w"]}">{merged_document_xml}</root>'
                        wrapped_tree = etree.fromstring(wrapped_xml.encode('utf-8'), parser=parser)
                        elements_to_add = list(wrapped_tree)

                    # 递归复制元素的函数
                    def deep_copy_element(source_elem, target_parent, target_ns):
                        """递归复制元素，保留所有属性和命名空间"""
                        if source_elem.tag.startswith('{'):
                            new_elem = etree.SubElement(target_parent, source_elem.tag)
                        else:
                            new_elem = etree.SubElement(target_parent, f"{{{target_ns}}}{source_elem.tag}")
                        
                        # 复制所有属性
                        for attr, value in source_elem.attrib.items():
                            new_elem.set(attr, value)
                        
                        # 复制文本内容
                        if source_elem.text is not None:
                            new_elem.text = source_elem.text
                        
                        # 递归复制子元素
                        for child in source_elem:
                            deep_copy_element(child, new_elem, target_ns)
                        
                        # 复制tail文本
                        if source_elem.tail is not None:
                            new_elem.tail = source_elem.tail
                        
                        return new_elem

                    # 清空body并添加所有元素
                    body = doc_tree.xpath('//w:body', namespaces=namespaces)[0]
                    for child in list(body):
                        body.remove(child)

                    # 按顺序添加所有元素
                    for element in elements_to_add:
                        deep_copy_element(element, body, namespaces['w'])

                    # 保存修改后的document.xml
                    new_doc_xml = etree.tostring(
                        doc_tree,
                        encoding='utf-8',
                        xml_declaration=True,
                        method='xml',
                        pretty_print=False
                    )
                    new_zip.writestr('word/document.xml', new_doc_xml)

                    # 写入合并后的styles.xml
                    if merged_styles_xml:
                        new_zip.writestr('word/styles.xml', merged_styles_xml.encode('utf-8'))
                    else:
                        # 如果没有合并的样式，使用模板的styles.xml
                        styles_xml = template_zip.read('word/styles.xml')
                        new_zip.writestr('word/styles.xml', styles_xml)

                # 将新ZIP文件写入磁盘
                output_buffer.seek(0)
                with open(temp_file, 'wb') as f:
                    f.write(output_buffer.getvalue())

            logger.info(f"文件保存成功: temp_file={temp_file}, tables_count={len(document_xml_list)}")

            # 添加后台任务清理临时文件
            background_tasks.add_task(cleanup_temp_file, temp_file)

            # 返回文件供下载
            logger.info(
                f"文件下载响应: filename=exported_tables.docx, tables_count={len(document_xml_list)}"
            )
            return FileResponse(
                temp_file,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename="exported_tables.docx",
            )

        except Exception as e:
            # 清理临时文件
            if temp_file and os.path.exists(temp_file):
                try:
                    os.unlink(temp_file)
                except:
                    pass
            logger.error(f"文件保存失败: error={e}", exc_info=True)
            raise HTTPException(
                status_code=HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"文件保存失败: {str(e)}",
            )

    except HTTPException:
        raise
    except Exception as e:
        # 清理临时文件
        if temp_file and os.path.exists(temp_file):
            try:
                os.unlink(temp_file)
            except:
                pass
        logger.error(f"导出失败: error={e}", exc_info=True)
        raise HTTPException(
            status_code=HTTP_500_INTERNAL_SERVER_ERROR, detail=f"导出失败: {str(e)}"
        )
