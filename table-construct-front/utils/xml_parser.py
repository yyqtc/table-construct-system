"""
XML解析和样式提取模块
实现完整的XML解析和样式提取逻辑
"""

import re
from typing import List, Dict, Any, Tuple
from docx.document import Document as DocumentType
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE


def extract_table_xml(table: Table) -> str:
    """
    提取表格的XML内容

    Args:
        table: python-docx Table对象

    Returns:
        表格的XML字符串
    """
    tbl_element = table._tbl
    return tbl_element.xml


def get_paragraphs_before_table(
    doc: DocumentType, table_index: int, count: int = 3
) -> List[str]:
    """
    获取表格前面的段落文本

    Args:
        doc: Document对象
        table_index: 表格在文档中的索引位置
        count: 需要获取的段落数量，默认3个

    Returns:
        段落文本列表
    """
    paragraphs = []

    tables = doc.tables
    if table_index >= len(tables):
        return paragraphs

    target_table_elem = tables[table_index]._tbl

    body = doc.element.body
    para_texts = []

    for elem in body:
        if elem == target_table_elem:
            break
        elif isinstance(elem, CT_P):
            para_text = ""
            for t_elem in elem.iter():
                if t_elem.tag.endswith("}t"):
                    if t_elem.text:
                        para_text += t_elem.text
            if para_text.strip():
                para_texts.append(para_text.strip())

    start_idx = max(0, len(para_texts) - count)
    paragraphs = para_texts[start_idx:]

    return paragraphs


def get_paragraphs_after_table(
    doc: DocumentType, table_index: int, count: int = 1
) -> List[str]:
    """
    获取表格后面的段落文本

    Args:
        doc: Document对象
        table_index: 表格在文档中的索引位置
        count: 需要获取的段落数量，默认1个

    Returns:
        段落文本列表
    """
    paragraphs = []

    tables = doc.tables
    if table_index >= len(tables):
        return paragraphs

    target_table_elem = tables[table_index]._tbl

    body = doc.element.body
    para_texts = []
    found_table = False

    for elem in body:
        if elem == target_table_elem:
            found_table = True
            continue
        elif found_table and isinstance(elem, CT_P):
            para_text = ""
            for t_elem in elem.iter():
                if t_elem.tag.endswith("}t"):
                    if t_elem.text:
                        para_text += t_elem.text
            if para_text.strip():
                para_texts.append(para_text.strip())
                if len(para_texts) >= count:
                    break

    paragraphs = para_texts

    return paragraphs


def _get_paragraph_objects_before_table(
    doc: DocumentType, table_index: int, count: int = 3
) -> List[Paragraph]:
    """
    获取表格前面的段落对象

    Args:
        doc: Document对象
        table_index: 表格在文档中的索引位置
        count: 需要获取的段落数量，默认3个

    Returns:
        段落对象列表
    """
    paragraphs = []

    tables = doc.tables
    if table_index >= len(tables):
        return paragraphs

    target_table_elem = tables[table_index]._tbl

    body = doc.element.body
    para_objs = []

    for elem in body:
        if elem == target_table_elem:
            break
        elif isinstance(elem, CT_P):
            para = Paragraph(elem, doc)
            para_text = para.text.strip()
            if para_text:
                para_objs.append(para)

    start_idx = max(0, len(para_objs) - count)
    paragraphs = para_objs[start_idx:]

    return paragraphs


def _get_paragraph_objects_after_table(
    doc: DocumentType, table_index: int, count: int = 1
) -> List[Paragraph]:
    """
    获取表格后面的段落对象

    Args:
        doc: Document对象
        table_index: 表格在文档中的索引位置
        count: 需要获取的段落数量，默认1个

    Returns:
        段落对象列表
    """
    paragraphs = []

    tables = doc.tables
    if table_index >= len(tables):
        return paragraphs

    target_table_elem = tables[table_index]._tbl

    body = doc.element.body
    para_objs = []
    found_table = False

    for elem in body:
        if elem == target_table_elem:
            found_table = True
            continue
        elif found_table and isinstance(elem, CT_P):
            para = Paragraph(elem, doc)
            para_text = para.text.strip()
            if para_text:
                para_objs.append(para)
                if len(para_objs) >= count:
                    break

    paragraphs = para_objs

    return paragraphs


def process_style_inheritance(
    table: Table, paragraphs_before: List[Paragraph], paragraphs_after: List[Paragraph]
) -> Dict[str, Any]:
    """
    处理样式继承逻辑

    Args:
        table: Table对象
        paragraphs_before: 表格前面的段落对象列表
        paragraphs_after: 表格后面的段落对象列表

    Returns:
        样式信息字典
    """
    style_metadata = {
        "table_styles": {},
        "paragraph_styles": {},
        "style_inheritance": {},
    }

    try:
        document = None

        if hasattr(table, "_parent") and table._parent is not None:
            if hasattr(table._parent, "document"):
                document = table._parent.document
            elif hasattr(table._parent, "part") and hasattr(
                table._parent.part, "document"
            ):
                document = table._parent.part.document

        if document is None and len(paragraphs_before) > 0:
            para = paragraphs_before[0]
            if hasattr(para, "_parent") and para._parent is not None:
                if hasattr(para._parent, "document"):
                    document = para._parent.document
                elif hasattr(para._parent, "part") and hasattr(
                    para._parent.part, "document"
                ):
                    document = para._parent.part.document

        if document is None and len(paragraphs_after) > 0:
            para = paragraphs_after[0]
            if hasattr(para, "_parent") and para._parent is not None:
                if hasattr(para._parent, "document"):
                    document = para._parent.document
                elif hasattr(para._parent, "part") and hasattr(
                    para._parent.part, "document"
                ):
                    document = para._parent.part.document

        if document is None or not hasattr(document, "styles"):
            return style_metadata

        def extract_style_xml(style):
            """提取样式的XML内容"""
            try:
                if hasattr(style, "_element"):
                    return style._element.xml
                elif hasattr(style, "element"):
                    return style.element.xml
            except:
                pass
            return ""

        def extract_based_on_from_xml(style_xml):
            """从XML中提取w:basedOn属性"""
            if not style_xml:
                return None
            match = re.search(r'w:basedOn[^=]*="([^"]*)"', style_xml, re.IGNORECASE)
            if match:
                return match.group(1)
            return None

        def get_based_on_chain(style, style_dict, visited=None):
            """递归获取样式继承链"""
            if visited is None:
                visited = set()

            style_id = style.style_id if hasattr(style, "style_id") else None
            if not style_id or style_id in visited:
                return []

            visited.add(style_id)
            chain = [style_id]

            based_on_id = None
            if hasattr(style, "based_on") and style.based_on is not None:
                based_on_id = (
                    style.based_on.style_id
                    if hasattr(style.based_on, "style_id")
                    else None
                )
            else:
                style_xml = extract_style_xml(style)
                based_on_id = extract_based_on_from_xml(style_xml)

            if based_on_id and based_on_id in style_dict:
                based_on_style = style_dict[based_on_id]
                based_on_chain = get_based_on_chain(based_on_style, style_dict, visited)
                chain.extend(based_on_chain)

            return chain

        def extract_style_info(style, all_styles_dict):
            """提取样式信息"""
            style_id = style.style_id if hasattr(style, "style_id") else None
            style_name = style.name if hasattr(style, "name") else None
            style_type = str(style.type) if hasattr(style, "type") else None
            style_xml = extract_style_xml(style)

            style_info = {
                "style_id": style_id,
                "name": style_name,
                "type": style_type,
                "xml": style_xml,
            }

            based_on_id = None
            if hasattr(style, "based_on") and style.based_on is not None:
                based_on_id = (
                    style.based_on.style_id
                    if hasattr(style.based_on, "style_id")
                    else None
                )
            else:
                based_on_id = extract_based_on_from_xml(style_xml)

            style_info["based_on"] = based_on_id
            style_info["inheritance_chain"] = (
                get_based_on_chain(style, all_styles_dict) if style_id else []
            )

            return style_info

        try:
            styles = document.styles
        except:
            return style_metadata

        all_styles_dict = {}
        for style in styles:
            try:
                style_id = style.style_id if hasattr(style, "style_id") else None
                if style_id:
                    all_styles_dict[style_id] = style
            except:
                continue

        table_styles_dict = {}
        for style in styles:
            try:
                if hasattr(style, "type") and style.type == WD_STYLE_TYPE.TABLE:
                    style_info = extract_style_info(style, all_styles_dict)
                    style_id = style_info["style_id"]
                    if style_id:
                        table_styles_dict[style_id] = style_info
            except:
                continue

        style_metadata["table_styles"] = table_styles_dict

        paragraph_styles_dict = {}
        all_paragraphs = paragraphs_before + paragraphs_after

        for para in all_paragraphs:
            try:
                if hasattr(para, "style") and para.style is not None:
                    style = para.style
                    style_id = style.style_id if hasattr(style, "style_id") else None

                    if style_id:
                        if style_id not in paragraph_styles_dict:
                            if style_id not in all_styles_dict:
                                all_styles_dict[style_id] = style
                            style_info = extract_style_info(style, all_styles_dict)
                            paragraph_styles_dict[style_id] = style_info
            except:
                continue

        style_metadata["paragraph_styles"] = paragraph_styles_dict

        inheritance_map = {}
        all_styles = list(table_styles_dict.values()) + list(
            paragraph_styles_dict.values()
        )

        for style_info in all_styles:
            style_id = style_info.get("style_id")
            if style_id:
                inheritance_map[style_id] = {
                    "based_on": style_info.get("based_on"),
                    "inheritance_chain": style_info.get("inheritance_chain", []),
                }

        style_metadata["style_inheritance"] = inheritance_map

        tbl_element = table._tbl
        tbl_xml = tbl_element.xml

        if tbl_xml:
            tbl_style_match = re.search(
                r'w:tblStyle[^=]*="([^"]*)"', tbl_xml, re.IGNORECASE
            )
            if tbl_style_match:
                tbl_style_id = tbl_style_match.group(1)
                style_metadata["table_applied_style"] = tbl_style_id
                if tbl_style_id in table_styles_dict:
                    style_metadata["table_applied_style_info"] = table_styles_dict[
                        tbl_style_id
                    ]

    except Exception as e:
        pass

    return style_metadata


def parse_document_xml(
    doc: DocumentType,
) -> Tuple[List[str], List[List[str]], List[List[str]], List[Dict[str, Any]]]:
    """
    解析文档中的所有表格，提取XML和样式信息

    Args:
        doc: Document对象

    Returns:
        Tuple包含:
        - tables: 表格XML字符串列表
        - before_paragraphs_list: 每个表格前面的段落文本列表的列表
        - after_paragraphs_list: 每个表格后面的段落文本列表的列表
        - used_styles: 每个表格使用的样式信息字典列表
    """
    tables = []
    before_paragraphs_list = []
    after_paragraphs_list = []
    used_styles = []

    try:
        doc_tables = doc.tables
        if not doc_tables:
            return tables, before_paragraphs_list, after_paragraphs_list, used_styles

        for table_index, table in enumerate(doc_tables):
            try:
                table_xml = extract_table_xml(table)
                tables.append(table_xml)

                paragraphs_before_text = get_paragraphs_before_table(
                    doc, table_index, count=3
                )
                paragraphs_after_text = get_paragraphs_after_table(
                    doc, table_index, count=1
                )

                before_paragraphs_list.append(paragraphs_before_text)
                after_paragraphs_list.append(paragraphs_after_text)

                paragraphs_before_objs = _get_paragraph_objects_before_table(
                    doc, table_index, count=3
                )
                paragraphs_after_objs = _get_paragraph_objects_after_table(
                    doc, table_index, count=1
                )

                style_info = process_style_inheritance(
                    table, paragraphs_before_objs, paragraphs_after_objs
                )
                used_styles.append(style_info)

            except Exception:
                continue

    except Exception:
        pass

    return tables, before_paragraphs_list, after_paragraphs_list, used_styles
